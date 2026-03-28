import os
import io
import asyncio
from docx import Document
from datetime import datetime
import google.generativeai as genai
from bot.config import GEMINI_API_KEY

genai.configure(api_key=GEMINI_API_KEY)
# We can use pro model for better text generation
model = genai.GenerativeModel('gemini-1.5-pro-latest')

async def generate_contract_text(lessor_name: str, lessee_name: str, item_name: str, price: str, deposit: str, dates: str) -> str:
    prompt = f"""
    Сгенерируй подробный и юридически грамотный текст Договора аренды оборудования между физическими лицами (или ИП/Компанией).
    Договор должен содержать следующие заполненные данные:
    - Арендодатель: {lessor_name}
    - Арендатор: {lessee_name}
    - Предмет аренды: {item_name}
    - Стоимость аренды: {price}
    - Залог: {deposit}
    - Сроки аренды: {dates}
    
    Договор должен включать: 
    1. Предмет договора
    2. Права и обязанности сторон
    3. Ответственность сторон (включая утерю и порчу оборудования)
    4. Порядок передачи и возврата
    5. Форс-мажор
    6. Место для подписей, адреса и реквизиты сторон
    
    Верни ТОЛЬКО чистый текст договора (включая заголовки), без символов экранирования кода или маркдауна (вроде ```), так как этот текст будет напрямую вставлен в .docx документ. Текст должен быть отформатирован абзацами.
    """
    try:
        response = await asyncio.to_thread(model.generate_content, prompt)
        # Strip code blocks if they exist
        text = response.text
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1]) if lines[-1].startswith("```") else "\n".join(lines[1:])
        return text
    except Exception as e:
        print(f"Contract AI generation error: {e}")
        return ""

async def generate_docx_contract(lessor_name: str, lessee_name: str, item_name: str, price: str, deposit: str, dates: str) -> io.BytesIO:
    text = await generate_contract_text(lessor_name, lessee_name, item_name, price, deposit, dates)
    
    doc = Document()
    doc.add_heading('Договор аренды оборудования', 0)
    
    if text:
        for p in text.split('\n'):
            if p.strip():
                if p.isupper() or len(p) < 40 and p.endswith(':'):
                    doc.add_heading(p.strip(), level=2)
                else:
                    doc.add_paragraph(p.strip())
    else:
        doc.add_paragraph("Произошла ошибка при обращении к ИИ сервису генерации договора.")
        
    doc.add_paragraph(f"\nДокумент сгенерирован автоматически AI-сервисом RentBot.\nДата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f
